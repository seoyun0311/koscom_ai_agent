# app_mcp/reports/fill_docx_template.py

from docx import Document
from typing import Dict, Any
import os
import logging

logger = logging.getLogger(__name__)


def fill_docx_template(
    template_path: str, 
    output_path: str, 
    context: Dict[str, Any]
) -> str:
    """
    DOCX 템플릿의 {{변수명}}을 context 값으로 치환하여 새로운 파일 생성.
    
    특징:
    - Run 단위로 치환 (서식 보존)
    - 표(table) 내부도 치환
    - 줄바금(\n) 처리 지원
    
    Args:
        template_path: 템플릿 DOCX 파일 경로
        output_path: 생성할 DOCX 파일 경로
        context: {"final_grade": "A", "period": "2025-10", ...}
    
    Returns:
        생성된 파일의 절대 경로
    
    Raises:
        FileNotFoundError: 템플릿 파일이 없을 때
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    logger.info("[fill_docx_template] Loading template: %s", template_path)
    doc = Document(template_path)

    # 문단(paragraph) 치환 - Run 단위로 (서식 보존)
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, context)

    # 표(table) 안 치환
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, context)

    # 출력 디렉토리 생성
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    # 저장
    doc.save(output_path)
    logger.info("[fill_docx_template] DOCX saved to: %s", output_path)
    
    return output_path


def _replace_in_paragraph(paragraph, context: Dict[str, Any]):
    """
    단락(paragraph) 내의 모든 Run에서 {{key}} 형태를 치환.
    
    Run 단위로 치환하므로 굵기/색상/폰트 등 서식이 보존됨.
    """
    for run in paragraph.runs:
        original_text = run.text
        
        # 각 context key에 대해 치환
        for key, value in context.items():
            placeholder = f"{{{{{key}}}}}"  # {{key}} 형태
            
            if placeholder in original_text:
                # 값을 문자열로 변환
                str_value = str(value)
                
                # 줄바꿈 처리
                if "\n" in str_value:
                    # 첫 줄만 run.text에 넣고
                    lines = str_value.split("\n")
                    original_text = original_text.replace(placeholder, lines[0])
                    
                    # 나머지 줄은 줄바꿈 추가
                    # (주의: 이 방법은 완벽하지 않음, 복잡한 경우 별도 처리 필요)
                    for line in lines[1:]:
                        run.add_break()
                        run.text += line
                else:
                    # 단순 치환
                    original_text = original_text.replace(placeholder, str_value)
        
        # 치환된 텍스트 적용
        run.text = original_text


def fill_docx_template_safe(
    template_path: str,
    output_path: str,
    context: Dict[str, Any]
) -> str:
    """
    에러 처리가 강화된 버전.
    
    템플릿이 없거나 치환 실패 시에도 빈 문서라도 생성.
    """
    try:
        return fill_docx_template(template_path, output_path, context)
    
    except FileNotFoundError:
        logger.warning(
            "[fill_docx_template_safe] Template not found, creating blank document"
        )
        
        # 템플릿 없으면 빈 문서 생성
        doc = Document()
        doc.add_heading("보고서 생성 실패", level=1)
        doc.add_paragraph(f"템플릿을 찾을 수 없습니다: {template_path}")
        doc.add_paragraph(f"기간: {context.get('period', 'N/A')}")
        doc.add_paragraph(f"최종 등급: {context.get('final_grade', 'N/A')}")
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    
    except Exception as e:
        logger.error(f"[fill_docx_template_safe] Failed: {e}", exc_info=True)
        raise