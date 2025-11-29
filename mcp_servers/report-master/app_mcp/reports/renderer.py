# app_mcp/reports/renderer.py
import json, os
from typing import Dict, List

from app_mcp.models import ComplianceFinding
from docx import Document  # ⬅ 워드 보고서 생성을 위한 추가 import


async def render_artifacts(
    report_id: str,
    period: str,
    narrative: str,
    findings: List[ComplianceFinding],
    inputs: dict,
    outdir: str
) -> Dict[str, str]:
    os.makedirs(outdir, exist_ok=True)

    # ---------- 1) HTML ----------
    html_path = os.path.join(outdir, f"{report_id}.html")
    html = f"""<!doctype html>
<html><head><meta charset="utf-8"><title>{report_id}</title>
<style>body{{font-family:ui-sans-serif,system-ui;max-width:920px;margin:40px auto;}}
h1{{margin-bottom:0}} .tag{{display:inline-block;padding:2px 8px;border-radius:8px;background:#eee;margin-left:8px}}
pre{{background:#f8f8f8;padding:12px;border-radius:8px;white-space:pre-wrap}}</style>
</head><body>
<h1>준수 보고서 <span class="tag">{period}</span></h1>
<h3>Report ID: {report_id}</h3>
<h2>요약</h2>
<pre>{narrative}</pre>
<h2>조항별 판단</h2>
<ul>
{''.join([f"<li><b>{f.article}</b> — <i>{f.status}</i><br/>{f.summary}</li>" for f in findings])}
</ul>
<h2>참고 원천(요약)</h2>
<pre>{json.dumps({k: json.loads(v.model_dump_json()) for k, v in inputs.items()}, ensure_ascii=False, indent=2)}</pre>
</body></html>"""
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)

    # ---------- 2) JSON ----------
    json_path = os.path.join(outdir, f"{report_id}.json")
    with open(json_path, "w", encoding="utf-8") as f:
        payload = {
            "report_id": report_id,
            "period": period,
            "narrative": narrative,
            "findings": [fi.model_dump() for fi in findings],
        }
        json.dump(payload, f, ensure_ascii=False, indent=2)

    # ---------- 3) DOCX (Word) ----------
    docx_path = os.path.join(outdir, f"{report_id}.docx")
    doc = Document()

    # 제목 및 기본 정보
    doc.add_heading(f"준수 보고서 ({period})", level=1)
    doc.add_paragraph(f"Report ID: {report_id}")

    # 1. 요약
    doc.add_heading("1. 요약", level=2)
    for line in narrative.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")  # 빈 줄도 유지

    # 2. 조항별 판단
    doc.add_heading("2. 조항별 판단", level=2)
    for fi in findings:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"[{fi.article}] {fi.status} - {fi.summary}")

    # 3. 참고 원천(요약)
    doc.add_heading("3. 참고 원천(요약)", level=2)
    # inputs dict를 JSON 문자열로 정리해서 그대로 붙이기
    try:
        summarized_inputs = {
            k: json.loads(v.model_dump_json())
            for k, v in inputs.items()
        }
    except Exception:
        summarized_inputs = {k: str(v) for k, v in inputs.items()}

    doc.add_paragraph(json.dumps(summarized_inputs, ensure_ascii=False, indent=2))

    doc.save(docx_path)

    # ▶ artifacts에 docx까지 포함해서 반환
    return {
        "html": html_path,
        "json": json_path,
        "docx": docx_path,
    }
